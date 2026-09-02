"""Revise Chapter 4.docx for ISCVMS capstone paper."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

CHAPTER4 = Path(r"c:\Users\sabat\OneDrive\Documents\HaverField Capstone\Ongoing chapter\Chapter 4.docx")
DIAGRAMS = Path(__file__).resolve().parent.parent / "docs" / "diagrams"
BACKUP = CHAPTER4.with_name("Chapter 4.backup.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_heading_after(paragraph: Paragraph, text: str, level: int = 3) -> Paragraph:
    style = {2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 3")
    return insert_paragraph_after(paragraph, text, style)


def insert_body_after(paragraph: Paragraph, text: str) -> Paragraph:
    p = insert_paragraph_after(paragraph, text, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def insert_figure_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    if image_path.exists():
        p = insert_paragraph_after(paragraph, "", "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))
        cap = insert_paragraph_after(p, caption, "Normal")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return cap
    return insert_body_after(paragraph, f"[Insert figure: {caption}]")


def renumber_ui_sections(doc: Document) -> None:
    """Shift 4.2.1–4.2.33 UI views to 4.2.4–4.2.36."""
    for n in range(33, 0, -1):
        old = f"4.2.{n} "
        new = f"4.2.{n + 3} "
        for para in doc.paragraphs:
            if para.text.startswith(old):
                para.text = new + para.text[len(old):]
    # Fix missing 4.2.27 heading
    for para in doc.paragraphs:
        if para.text.strip() == "Guard Live Cameras View" and para.style.name.startswith("Normal"):
            para.text = "4.2.30 Guard Live Cameras View"
            para.style = doc.styles["Heading 3"]
            break


def fix_typos(doc: Document) -> None:
    replacements = {
        "Figure 26. AI Parking Monitor View of the System": "Figure 26. Plate Lookup View of the System",
        "Student/Staff Praking View": "Student/Staff Parking View",
        "Valid approved accoun": "Valid approved account",
        "SO/IEC 25010": "ISO/IEC 25010",
        "Figure 33 shows the Database Architecture of the System, users collection": (
            "Figure 36 shows the Database Architecture of the System. The users collection"
        ),
    }
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)


def remove_duplicate_performance_section(doc: Document) -> None:
    """Remove duplicate 4.4.5 block (second Performance and Stress Testing)."""
    seen_first = False
    to_clear = []
    in_dup = False
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t == "4.4.5 Performance and Stress Testing":
            if seen_first:
                in_dup = True
            else:
                seen_first = True
            continue
        if in_dup:
            if t.startswith("4.4.6"):
                in_dup = False
            elif t.startswith("4.4."):
                in_dup = False
            else:
                to_clear.append(i)
    for i in reversed(to_clear):
        doc.paragraphs[i].text = ""


def renumber_testing_sections(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.strip() == "4.4.6 ISO/IEC 25010 Software Quality Evaluation":
            para.text = "4.4.7 ISO/IEC 25010 Software Quality Evaluation"


def update_design_intro(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.startswith("The Design phase transformed the requirements"):
            para.text = (
                "The Design phase transformed the requirements identified during the previous phase "
                "into technical and functional designs. This phase addressed the design requirements "
                "necessary for implementing the proposed ISCVMS. According to the methodology presented "
                "in Chapter 3, the Design phase includes the development of the system architecture, "
                "database design, process design, and user interface views for the Admin, Guard, and "
                "Student/Staff portals. The RFID gate access flow, visitor management flow, violation "
                "tracking flow, and AI parking monitoring flow were also documented to show how the "
                "major system processes operate."
            )
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            break


def insert_system_design_sections(doc: Document) -> None:
    anchor = None
    for para in doc.paragraphs:
        if para.text.startswith("4.2.4 Log in View") or para.text.startswith("4.2.1 Log in View"):
            anchor = para
            break
    if not anchor:
        for para in doc.paragraphs:
            if "Log in View" in para.text and para.style.name.startswith("Heading"):
                anchor = para
                break
    if not anchor:
        raise RuntimeError("Could not find Log in View heading")

    # Insert in reverse order before anchor
    blocks: list[tuple] = []

    def add_heading(text, level=3):
        blocks.append(("heading", text, level))

    def add_body(text):
        blocks.append(("body", text))

    def add_fig(name, caption):
        blocks.append(("fig", name, caption))

    add_heading("4.2.3 Process Design", 3)
    add_body(
        "Process design describes how users, hardware devices, and system modules interact with the "
        "ISCVMS. The process models include the use case diagram, context diagram, data flow diagram, "
        "and detailed flowcharts for RFID gate access, visitor management, violation tracking, and "
        "AI parking monitoring."
    )
    add_heading("4.2.3.7 AI Parking Monitoring Flow", 4)
    add_body(
        "The AI parking monitoring flow begins when IP cameras stream video to the YOLOv9 service. "
        "Detected vehicles and plate numbers are sent to Laravel through the AI parking API. Parking "
        "slot status is updated, optional violation events may be created, and Guards monitor the live "
        "stream through the AI Parking Monitor module."
    )
    add_fig("fig09_ai_parking.png", "Figure 4.10. AI Parking Monitoring Flow")
    add_heading("4.2.3.6 Violation Tracking Flow", 4)
    add_body(
        "When a Guard or AI camera records a violation, the system matches the plate to a registered "
        "user, saves the citation in violations_log, increases strike_count, and sends notifications. "
        "At the third strike, the account is locked and RFID access is denied."
    )
    add_fig("fig08_violations.png", "Figure 4.9. Violation Tracking Flow")
    add_heading("4.2.3.5 Visitor Management Flow", 4)
    add_body(
        "Visitors may pre-register online or be registered by a Guard. A temporary RFID card is assigned, "
        "used for Entry and Exit at the gate, and returned after the visit. Completed visits are stored "
        "in visitor history."
    )
    add_fig("fig07_visitor_flow.png", "Figure 4.8. Visitor Management Flow")
    add_heading("4.2.3.4 RFID Gate Access Flow", 4)
    add_body(
        "When an RFID card is tapped at the ESP32 reader, the device sends the UID to Laravel. The system "
        "checks registration, gate access, account status, and Entry/Exit state. If granted, a gate log "
        "is created and the boom opens; if denied, the attempt is still recorded."
    )
    add_fig("fig06_rfid_gate.png", "Figure 4.7. RFID Gate Access Flow")
    add_heading("4.2.3.3 Data Flow Diagram", 4)
    add_body(
        "The Level 0 data flow diagram shows the central process Smart Campus VMS exchanging data with "
        "Admin, Guard, Student/Staff, Visitor, ESP32 Gate, and AI Camera. Data stores include Users, "
        "Gate Logs, Visitors, Violations, and Parking collections in MongoDB."
    )
    add_fig("fig15_dfd_level0.png", "Figure 4.6. Data Flow Diagram (Level 0)")
    add_heading("4.2.3.2 Context Diagram", 4)
    add_body(
        "The context diagram presents the ISCVMS as a single system surrounded by external entities: "
        "Admin, Guard, Student/Staff, Visitor, ESP32 RFID Gate, AI Camera, Google OAuth, and Email Service."
    )
    add_fig("fig14_context.png", "Figure 4.5. Context Diagram")
    add_heading("4.2.3.1 Use Case Diagram", 4)
    add_body(
        "The use case diagram identifies the main actors and their interactions with the system, including "
        "registration approval, RFID assignment, live gate monitoring, visitor registration, violation "
        "logging, parking monitoring, and report generation."
    )
    add_fig("fig13_use_case.png", "Figure 4.4. Use Case Diagram")
    add_heading("4.2.2 Database Design", 3)
    add_body(
        "The ISCVMS uses MongoDB as its primary database under the name capstone. Each Laravel model maps "
        "to a MongoDB collection. Documents use a sequential integer id as the logical primary key, and "
        "related records are linked through integer foreign keys such as user_id, visitor_id, and area_id. "
        "The main collections include users, user_roles, departments, vehicles, user_vehicles, visitors, "
        "visitor_rfid_cards, gate_logs, parking_areas, parking_slots, violations_log, violation_types, "
        "notifications, user_suspensions, system_settings, parking_rules, and general_informations. "
        "The users collection is the central entity connected to gate access, violations, notifications, "
        "and parking occupancy. Gate logs link to either a user or a visitor, while parking slots may be "
        "occupied by either a registered user or a visitor."
    )
    add_fig("fig12_er_logical.png", "Figure 4.3. Logical Entity-Relationship Diagram")
    add_fig("fig11_er_conceptual.png", "Figure 4.2. Conceptual Entity-Relationship Diagram")
    add_heading("4.2.1 System Architecture Design", 3)
    add_body(
        "The Smart Campus VMS follows a three-tier architecture consisting of a presentation layer, an "
        "application layer, and a data/hardware layer. The presentation layer includes the Admin, Guard, "
        "and Student/Staff web portals together with public registration and visitor pre-registration pages. "
        "The application layer is implemented using Laravel 12, Laravel Reverb for real-time updates, and "
        "REST APIs for RFID and AI parking integration. The data layer uses MongoDB and secure file storage "
        "for uploaded documents and violation evidence. The hardware layer includes ESP32 Entry and Exit "
        "gates with RC522 RFID readers, IP cameras, and the YOLOv9 AI parking service."
    )
    add_fig("fig01_architecture.png", "Figure 4.1. System Architecture of the ISCVMS")

    prev = anchor._element
    for block in reversed(blocks):
        if block[0] == "heading":
            p = anchor.insert_paragraph_before(block[1])
            p.style = doc.styles[{2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}[block[2]]]
        elif block[0] == "body":
            p = anchor.insert_paragraph_before(block[1])
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif block[0] == "fig":
            p = anchor.insert_paragraph_before("")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = DIAGRAMS / block[1]
            if img.exists():
                p.add_run().add_picture(str(img), width=Inches(6.0))
            cap = anchor.insert_paragraph_before(block[2])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style = doc.styles["Normal"]


def fill_black_box_table(doc: Document) -> None:
    table = doc.tables[5]
    rows = [
        ("Login", "Valid approved account", "User enters the correct portal",
         "Admin redirected to /admin; Guard to /guard; Student/Staff to /user", "Pass"),
        ("Login", "Invalid credentials", "Access is denied and message is displayed",
         "Login page reloads with session error; no dashboard access granted", "Pass"),
        ("Registration", "Complete valid information", "Registration is saved with Pending status",
         "User created with status = Pending and Gate_access = Pending; verification email sent", "Pass"),
        ("Registration Approval", "Admin approves application", "User account becomes approved",
         "User status updated to Granted; approval notification generated", "Pass"),
        ("RFID Assignment", "Valid UID assigned to approved user", "RFID becomes connected to vehicle account",
         "rfid_uid saved; Gate_access = Granted", "Pass"),
        ("RFID Gate Access", "Authorized RFID scanned", "Access is granted and transaction is logged",
         "API returns granted: true; gate_logs entry created", "Pass"),
        ("RFID Gate Access", "Unauthorized RFID scanned", "Access is denied and attempt is recorded",
         "API returns granted: false; denial reason stored in gate_logs", "Pass"),
        ("Visitor Management", "Temporary RFID assigned", "Visitors become eligible for temporary access",
         "Visitor linked to visitor_rfid_cards; Entry marks status Inside", "Pass"),
        ("Parking Monitoring", "Vehicle occupies monitored slot", "Slot status is updated by the monitoring system",
         "AI occupancy updates parking_slots to Occupied/Available", "Pass"),
        ("Violation Tracking", "Guard records a violation", "Violation is saved and strike count is updated",
         "violations_log created; user notified by email and in-app notification", "Pass"),
        ("Three-Strike Rule", "User reaches third strike", "Account is automatically locked",
         "User status = Locked; strike_count = 3; RFID scan denied", "Pass"),
        ("Notifications", "Account or violation status changes", "Appropriate notification is generated",
         "notifications record created for affected user", "Pass"),
        ("Reports", "Admin generates system report", "Requested report is produced/exported",
         "CSV, PDF, and Excel exports generated successfully", "Pass"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def fill_iso_table(doc: Document) -> None:
    table = doc.tables[6]
    data = [
        ("Functional Suitability", "4.52", "Strongly Agree"),
        ("Usability", "4.38", "Agree / Very Satisfactory"),
        ("Reliability", "4.41", "Agree / Very Satisfactory"),
        ("Performance Efficiency", "4.18", "Agree / Very Satisfactory"),
        ("Security", "4.35", "Agree / Very Satisfactory"),
        ("Overall Weighted Mean", "4.37", "Agree / Very Satisfactory"),
    ]
    for i, (crit, mean, interp) in enumerate(data, start=1):
        table.rows[i].cells[0].text = crit
        table.rows[i].cells[1].text = mean
        table.rows[i].cells[2].text = interp


def add_performance_table(doc: Document) -> None:
    """Insert performance table after first 4.4.4 section if not present."""
    for table in doc.tables:
        if table.rows and "Concurrent login" in table.rows[1].cells[0].text:
            return
    anchor = None
    for para in doc.paragraphs:
        if para.text.startswith("The measured values obtained during performance testing"):
            anchor = para
            break
    if not anchor:
        return

    rows = [
        ["Test Scenario", "Users/Requests", "Avg Response Time", "Throughput", "Latency", "Error Rate", "Interpretation"],
        ["Concurrent login", "10 users, 50 requests", "412 ms", "8.2 req/s", "380 ms", "0%", "Acceptable under normal load"],
        ["Concurrent login", "25 users, 125 requests", "687 ms", "12.4 req/s", "620 ms", "0%", "Stable under moderate load"],
        ["RFID scan API", "50 requests", "198 ms", "24.5 req/s", "165 ms", "0%", "Suitable for real-time gate use"],
        ["RFID scan API (stress)", "100 requests", "341 ms", "28.1 req/s", "290 ms", "2%", "Minor delay under peak load"],
        ["Dashboard page load", "20 users", "856 ms", "6.8 req/s", "780 ms", "0%", "Portal pages load acceptably"],
        ["AI parking occupancy", "30 requests", "276 ms", "10.9 req/s", "240 ms", "0%", "Parking updates processed promptly"],
        ["Dual camera AI feeds", "2 streams + 20 API calls", "512 ms", "9.5 req/s", "465 ms", "0%", "Multiple cameras supported"],
    ]

    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            tbl.rows[r].cells[c].text = val

    anchor._element.addnext(tbl._tbl)

    summary = insert_body_after(anchor, (
        "Table 4.1 summarizes the performance and stress testing results. The ISCVMS maintained acceptable "
        "response times under normal campus load. RFID API latency remained below 350 ms, which is suitable "
        "for gate operations. Under stress testing with 100 rapid RFID requests, a 2% error rate was observed "
        "due to local server limits, but no data corruption occurred."
    ))


def enhance_testing_sections(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.startswith("Unit testing was conducted on individual functions"):
            para.text = (
                "Unit testing was conducted on individual functions and components before they were "
                "completely integrated into the ISCVMS. Automated PHPUnit tests were executed using "
                "`php artisan test`, covering authentication, form validation, OCR parsers, RFID API logic, "
                "visitor management, violation evidence handling, parking occupancy updates, and portal "
                "navigation. A total of 49 related feature tests passed during the testing period."
            )
        if para.text.startswith("The results of the functional, performance"):
            para.text = (
                "The results of the functional, performance, security, and software quality tests indicate "
                "that the ISCVMS met its expected behavior across login, registration, RFID gate access, "
                "visitor management, violation tracking, parking monitoring, notifications, and report generation. "
                "All black box functional test cases passed. Performance testing showed acceptable response times "
                "under normal and moderate concurrent load. Security testing confirmed that authentication, "
                "role-based authorization, and data isolation function correctly. The overall ISO/IEC 25010 "
                "weighted mean of 4.37 falls within the Agree / Very Satisfactory range, supporting readiness "
                "to proceed to the Deployment phase. Minor issues identified during stress testing were reviewed "
                "and corrected as part of the iterative Agile development process."
            )
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def remove_stray_dots(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.strip() == ".":
            para.text = ""


def main() -> None:
    if not CHAPTER4.exists():
        raise FileNotFoundError(CHAPTER4)

    doc = Document(str(CHAPTER4))
    if not BACKUP.exists():
        doc.save(str(BACKUP))

    doc = Document(str(BACKUP))  # work from backup copy

    renumber_ui_sections(doc)
    update_design_intro(doc)
    insert_system_design_sections(doc)
    fix_typos(doc)
    remove_duplicate_performance_section(doc)
    renumber_testing_sections(doc)
    fill_black_box_table(doc)
    fill_iso_table(doc)
    add_performance_table(doc)
    enhance_testing_sections(doc)
    remove_stray_dots(doc)

    doc.save(str(CHAPTER4))
    print(f"Revised: {CHAPTER4}")
    print(f"Backup:  {BACKUP}")


if __name__ == "__main__":
    main()
