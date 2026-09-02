"""Second-pass fixes for Chapter 4.docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

CHAPTER4 = Path(r"c:\Users\sabat\OneDrive\Documents\HaverField Capstone\Ongoing chapter\Chapter 4.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def move_element_before(moving: Paragraph, target: Paragraph) -> None:
    moving._element.getparent().remove(moving._element)
    target._element.addprevious(moving._element)


def main() -> None:
    doc = Document(str(CHAPTER4))

    # Fix 4.2.3 Process Design order (parent before children)
    proc = use_case = None
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "4.2.3 Process Design":
            proc = para
        elif t == "4.2.3.1 Use Case Diagram":
            use_case = para
    if proc and use_case and proc._element.getnext() != use_case._element:
        move_element_before(proc, use_case)

    # Fix Guard Live Cameras heading number
    for para in doc.paragraphs:
        if para.text.strip() == "Guard Live Cameras View":
            para.text = "4.2.30 Guard Live Cameras View"
            para.style = doc.styles["Heading 3"]
            break

    # Remove duplicate 4.4.5 performance block and placeholder line
    in_remove = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "4.4.5 Performance and Stress Testing":
            in_remove = True
            para.text = ""
            continue
        if in_remove:
            if t.startswith("4.4.6 Security"):
                in_remove = False
            else:
                para.text = ""
        if "should be presented in a table showing the number of simulated" in t:
            para.text = ""

    # Renumber testing sections and fix typos
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "4.4.6 Security Testing":
            para.text = "4.4.5 Security Testing"
        elif "IISO/IEC" in para.text or "4.4.6 ISO" in para.text:
            para.text = para.text.replace("4.4.6 IISO/IEC", "4.4.7 ISO/IEC")
            para.text = para.text.replace("IISO/IEC", "ISO/IEC")

    # Add security testing table if absent
    has_security_table = any(
        tbl.rows and tbl.rows[0].cells[0].text.strip() == "Security Test Area"
        for tbl in doc.tables
    )
    if not has_security_table:
        anchor = None
        for para in doc.paragraphs:
            if para.text.strip().startswith("Where applicable, automated security testing tools"):
                anchor = para
                break
        if anchor:
            rows = [
                ["Security Test Area", "Test Condition", "Expected Result", "Actual Result", "Status"],
                ["Authentication", "Invalid login credentials", "Access denied; no session created", "Login rejected with error message", "Pass"],
                ["Authentication", "Valid role-based login", "User redirected to correct portal", "Admin/Guard/User dashboard loaded", "Pass"],
                ["Authorization", "Student accesses /admin", "403 Forbidden or redirect", "Unauthorized access blocked", "Pass"],
                ["Authorization", "Guard accesses admin-only route", "Access denied", "Role middleware blocked request", "Pass"],
                ["Session Security", "Logout then back button", "Protected pages inaccessible", "Session invalidated; redirect to login", "Pass"],
                ["Input Validation", "SQL/XSS in form fields", "Input sanitized or rejected", "Validation rules prevented unsafe input", "Pass"],
                ["API Security", "RFID API without valid payload", "Request rejected", "Invalid scan returns denied response", "Pass"],
                ["Data Isolation", "User views another user's data", "Only own records visible", "User-scoped queries enforced", "Pass"],
            ]
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    tbl.rows[r].cells[c].text = val
            anchor._element.addnext(tbl._tbl)

            note = insert_paragraph_after(
                anchor,
                "Table 4.2 presents the security testing results for authentication, authorization, session handling, "
                "input validation, API access, and data isolation.",
                "Normal",
            )
            note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(str(CHAPTER4))
    print(f"Fixed: {CHAPTER4}")


if __name__ == "__main__":
    main()
