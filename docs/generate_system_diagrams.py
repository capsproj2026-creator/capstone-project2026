"""Generate Smart Campus VMS dashboard diagrams, ER diagrams, and a Word document."""

from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent
DIAG_DIR = OUT_DIR / "diagrams"
DOCX_PATH = OUT_DIR / "Smart-Campus-VMS-Dashboards.docx"

NAVY = (15, 45, 90)
NAVY_SOFT = (232, 239, 248)
GOLD = (196, 149, 58)
TEAL = (15, 118, 110)
GREEN = (22, 101, 52)
RED = (153, 27, 27)
ORANGE = (146, 64, 14)
PURPLE = (88, 28, 135)
GRAY = (71, 85, 105)
WHITE = (255, 255, 255)
BG = (248, 250, 252)
LINE = (51, 65, 85)
PK_BG = (254, 249, 195)
FK_BG = (219, 234, 254)


def font(size: int, bold: bool = False):
    names = ["segoeui.ttf", "SegoeUI.ttf", "arial.ttf", "Arial.ttf", "calibri.ttf"]
    bold_names = ["segoeuib.ttf", "SegoeUI-Bold.ttf", "arialbd.ttf", "ArialBd.ttf"]
    windir = Path(r"C:\Windows\Fonts")
    for name in bold_names if bold else names:
        path = windir / name
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw, text: str, fnt, max_w: int) -> list[str]:
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if draw.textlength(trial, font=fnt) <= max_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [text]


def new_canvas(w: int, h: int, title: str):
    img = Image.new("RGB", (w, h), BG)
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 72), fill=NAVY)
    d.text((28, 22), title, fill=WHITE, font=font(24, True))
    d.rectangle((0, 72, w, 78), fill=GOLD)
    return img, d


def rounded(d, xy, fill, outline, r=14, width=2):
    d.rounded_rectangle(xy, radius=r, fill=fill, outline=outline, width=width)


def box(d, x, y, w, h, text, fill, outline=NAVY, text_color=NAVY):
    rounded(d, (x, y, x + w, y + h), fill, outline, r=12)
    fnt = font(15, True)
    lines = wrap(d, text, fnt, w - 20)
    total_h = len(lines) * 20
    ty = y + (h - total_h) // 2
    for i, line in enumerate(lines):
        tw = d.textlength(line, font=fnt)
        d.text((x + (w - tw) / 2, ty + i * 20), line, fill=text_color, font=fnt)
    return (x + w / 2, y, x + w / 2, y + h, x, y + h / 2, x + w, y + h / 2)


def diamond(d, cx, cy, w, h, text, fill=(255, 247, 237), outline=ORANGE):
    pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
    d.polygon(pts, fill=fill, outline=outline)
    fnt = font(14, True)
    lines = wrap(d, text, fnt, w - 36)
    total_h = len(lines) * 18
    ty = cy - total_h / 2
    for i, line in enumerate(lines):
        tw = d.textlength(line, font=fnt)
        d.text((cx - tw / 2, ty + i * 18), line, fill=outline, font=fnt)
    return cx, cy - h / 2, cx, cy + h / 2, cx - w / 2, cy, cx + w / 2, cy


def arrow(d, x1, y1, x2, y2, color=LINE, label=None):
    d.line((x1, y1, x2, y2), fill=color, width=3)
    ang = math.atan2(y2 - y1, x2 - x1)
    s = 11
    d.polygon(
        [
            (x2, y2),
            (x2 - s * math.cos(ang - 0.45), y2 - s * math.sin(ang - 0.45)),
            (x2 - s * math.cos(ang + 0.45), y2 - s * math.sin(ang + 0.45)),
        ],
        fill=color,
    )
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        fnt = font(12, True)
        tw = d.textlength(label, font=fnt)
        d.rectangle((mx - tw / 2 - 5, my - 11, mx + tw / 2 + 5, my + 11), fill=BG)
        d.text((mx - tw / 2, my - 8), label, fill=color, font=fnt)


def save(img: Image.Image, name: str) -> Path:
    DIAG_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAG_DIR / name
    img.save(path, "PNG", optimize=True)
    return path


# --- ER helpers ---

def entity(d, x, y, w, title, fields, header=NAVY):
    row_h = 20
    h = 34 + len(fields) * row_h + 6
    d.rounded_rectangle((x, y, x + w, y + h), radius=8, fill=WHITE, outline=header, width=2)
    d.rectangle((x, y, x + w, y + 32), fill=header)
    d.text((x + 10, y + 7), title, fill=WHITE, font=font(13, True))
    yy = y + 34
    fnt = font(12)
    for kind, name in fields:
        bg = PK_BG if kind == "PK" else FK_BG if kind == "FK" else WHITE
        d.rectangle((x + 1, yy, x + w - 1, yy + row_h), fill=bg)
        tag = {"PK": "PK", "FK": "FK", "A": "  "}.get(kind, "  ")
        d.text((x + 8, yy + 2), f"{tag}  {name}", fill=NAVY, font=fnt)
        yy += row_h
    return {"x": x, "y": y, "w": w, "h": h, "l": x, "r": x + w, "t": y, "b": y + h, "cx": x + w / 2, "cy": y + h / 2}


def ports(e, side: str):
    if side == "n":
        return e["cx"], e["t"]
    if side == "s":
        return e["cx"], e["b"]
    if side == "w":
        return e["l"], e["cy"]
    return e["r"], e["cy"]


def crow(d, x, y, toward, many=True, optional=False, color=LINE):
    """Draw cardinality at (x,y). toward is angle of the line leaving the entity."""
    ang = toward
    ux, uy = math.cos(ang), math.sin(ang)
    px, py = -uy, ux
    if optional:
        d.ellipse((x - 5 + ux * 6, y - 5 + uy * 6, x + 5 + ux * 6, y + 5 + uy * 6), outline=color, width=2)
        base = 14
    else:
        d.line((x + px * 8, y + py * 8, x - px * 8, y - py * 8), fill=color, width=3)
        base = 8
    if many:
        tipx, tipy = x + ux * (base + 12), y + uy * (base + 12)
        d.line((x + ux * base, y + uy * base, tipx + px * 8, tipy + py * 8), fill=color, width=3)
        d.line((x + ux * base, y + uy * base, tipx - px * 8, tipy - py * 8), fill=color, width=3)
        d.line((x + ux * base, y + uy * base, tipx, tipy), fill=color, width=3)


def relate(d, a, sa, b, sb, a_card="1", b_card="N", label="", color=LINE):
    x1, y1 = ports(a, sa)
    x2, y2 = ports(b, sb)
    d.line((x1, y1, x2, y2), fill=color, width=2)
    ang_ab = math.atan2(y2 - y1, x2 - x1)
    crow(d, x1, y1, ang_ab, many=a_card in ("N", "0..N"), optional=a_card.startswith("0"), color=color)
    crow(d, x2, y2, ang_ab + math.pi, many=b_card in ("N", "0..N"), optional=b_card.startswith("0"), color=color)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        fnt = font(11, True)
        tw = d.textlength(label, font=fnt)
        d.rectangle((mx - tw / 2 - 4, my - 10, mx + tw / 2 + 4, my + 10), fill=BG)
        d.text((mx - tw / 2, my - 8), label, fill=color, font=fnt)


def fig_architecture():
    img, d = new_canvas(1400, 980, "Figure 1. Smart Campus VMS — System Architecture")
    layers = [
        (110, [(40, "ESP32 Entry · RC522 + Servo GPIO 14"), (370, "ESP32 Exit · RC522 only"),
               (700, "IP / Tapo cameras"), (1030, "Owner + visitor RFID cards")]),
        (340, [(40, "Laravel 12 · port 8000 · MongoDB"), (370, "RFID APIs /scan + /heartbeat"),
               (700, "YOLOv9 AI parking · port 8090"), (1030, "Laravel Reverb live events")]),
        (570, [(120, "Admin portal /admin"), (520, "Guard portal /guard"), (920, "Student/Staff portal /user")]),
        (800, [(220, "MongoDB collections (this ER)"), (780, "File storage · IDs, OR/CR, evidence")]),
    ]
    colors = [(NAVY_SOFT, NAVY), ((224, 242, 254), (3, 105, 161)), ((237, 233, 254), PURPLE), ((254, 243, 199), ORANGE)]
    ws = [300, 300, 360, 400]
    for i, (y, boxes) in enumerate(layers):
        fill, outline = colors[i]
        w = ws[i]
        for x, text in boxes:
            box(d, x, y, w, 90, text, fill, outline)
    for x in (190, 520, 850, 1180):
        arrow(d, x, 200, x, 330)
    d.text((40, 930), "Laravel stores all operational data in MongoDB. ESP32 opens the boom only when granted = true.", fill=GRAY, font=font(15))
    return save(img, "fig01_architecture.png")


def fig_login():
    img, d = new_canvas(1400, 1100, "Figure 2. Login, verification, and role routing")
    b1 = box(d, 520, 110, 360, 70, "User opens /login", (219, 234, 254), (3, 105, 161))
    b2 = box(d, 520, 230, 360, 70, "Submit email + password", NAVY_SOFT)
    dia = diamond(d, 700, 390, 280, 110, "Valid credentials?")
    b3 = box(d, 40, 470, 280, 80, "Invalid login error", (254, 226, 226), RED, RED)
    dia2 = diamond(d, 700, 560, 280, 110, "Email verified?")
    b4 = box(d, 1040, 520, 300, 80, "Email verification page", (254, 243, 199), ORANGE, ORANGE)
    dia3 = diamond(d, 700, 740, 300, 120, "Granted and not locked?")
    b5 = box(d, 40, 820, 300, 80, "Block pending / denied / 3-strike lock", (254, 226, 226), RED, RED)
    box(d, 160, 960, 240, 70, "Admin → /admin", (219, 234, 254), NAVY)
    box(d, 580, 960, 240, 70, "Guard → /guard", (204, 251, 241), TEAL)
    box(d, 1000, 960, 280, 70, "Student/Staff → /user", (237, 233, 254), PURPLE)
    arrow(d, b1[2], b1[3], b2[0], b2[1])
    arrow(d, b2[2], b2[3], dia[0], dia[1])
    arrow(d, dia[4], dia[5], b3[6], b3[5], RED, "No")
    arrow(d, dia[2], dia[3], dia2[0], dia2[1], GREEN, "Yes")
    arrow(d, dia2[6], dia2[7], b4[4], b4[5], ORANGE, "No")
    arrow(d, dia2[2], dia2[3], dia3[0], dia3[1], GREEN, "Yes")
    arrow(d, dia3[4], dia3[5], b5[6], b5[5], RED, "No")
    arrow(d, dia3[2], dia3[3], 700, 960, GREEN, "Yes")
    return save(img, "fig02_login_roles.png")


def fig_admin_map():
    img, d = new_canvas(1500, 1080, "Figure 3. Admin dashboard map  |  /admin")
    box(d, 500, 110, 500, 90, "Admin Dashboard — KPIs, charts, 2nd-strike alerts", (219, 234, 254), NAVY)
    modules = [
        (40, 260, "Registrations", "Approve or decline Student/Staff accounts"),
        (400, 260, "RFID Assignment", "Bind RFID UID and grant gate access"),
        (760, 260, "User Management", "Users, documents, strikes, status"),
        (1120, 260, "Settings", "Rules, violation types, admins, parking"),
        (40, 460, "Violations", "Review citations and evidence"),
        (400, 460, "Access Logs", "Live entry/exit history"),
        (760, 460, "Parking", "Areas, slots, zone access"),
        (1120, 460, "Live Cameras", "AI / CCTV feeds"),
        (40, 660, "Reports", "Export PDF / Excel"),
        (400, 660, "Registered Visitors", "Guests currently on campus"),
        (760, 660, "Visitor History", "Completed and expired visits"),
        (1120, 660, "Create Guard", "Register a Guard account"),
    ]
    for x, y, title, desc in modules:
        rounded(d, (x, y, x + 330, y + 140), WHITE, NAVY)
        d.rectangle((x, y, x + 330, y + 38), fill=NAVY)
        d.text((x + 14, y + 10), title, fill=WHITE, font=font(16, True))
        fnt = font(14)
        for i, line in enumerate(wrap(d, desc, fnt, 300)):
            d.text((x + 14, y + 54 + i * 20), line, fill=GRAY, font=fnt)
    widgets = ["Total / active / suspended users", "Active violations (3-strike)", "Today entries & exits",
               "Parking occupied / available", "Visitor counts", "Weekly trends", "Violation type chart",
               "2nd-strike alert list", "Recent violations", "Quick links"]
    for i, w in enumerate(widgets):
        box(d, 40 + (i % 5) * 290, 860 + (i // 5) * 70, 270, 56, w, NAVY_SOFT, NAVY)
    return save(img, "fig03_admin_dashboard.png")


def fig_registration():
    img, d = new_canvas(1400, 1180, "Figure 4. Registration → approval → RFID assignment")
    steps = [(520, 110, "Student/Staff fills /register"), (520, 230, "Email verification"),
             (520, 350, "Status = Pending"), (520, 470, "Admin Registrations: Approve or Decline")]
    for i, (x, y, t) in enumerate(steps):
        box(d, x, y, 380, 80, t, NAVY_SOFT)
        if i:
            arrow(d, 710, y - 40, 710, y)
    dia = diamond(d, 710, 640, 300, 120, "Approved?")
    arrow(d, 710, 550, dia[0], dia[1])
    no = box(d, 40, 700, 320, 90, "Denied — no portal access", (254, 226, 226), RED, RED)
    yes = box(d, 980, 600, 360, 90, "Granted portal. RFID still off.", (220, 252, 231), GREEN, GREEN)
    arrow(d, dia[4], dia[5], no[6], no[5], RED, "Decline")
    arrow(d, dia[6], dia[7], yes[4], yes[5], GREEN, "Approve")
    rfid = box(d, 500, 820, 420, 90, "Admin RFID Assignment: UID + Gate_access", (219, 234, 254), NAVY)
    arrow(d, 1160, 690, 1160, 865)
    arrow(d, 1160, 865, rfid[6], rfid[5])
    gate = box(d, 500, 980, 420, 90, "User taps RFID at ESP32 Entry/Exit", (204, 251, 241), TEAL)
    arrow(d, rfid[2], rfid[3], gate[0], gate[1])
    return save(img, "fig04_registration_rfid.png")


def fig_guard_map():
    img, d = new_canvas(1500, 1120, "Figure 5. Guard dashboard map  |  /guard")
    box(d, 470, 110, 560, 90, "Guard Dashboard — operations overview + shortcuts", (204, 251, 241), TEAL)
    modules = [
        (40, 250, "Live Gate Monitor", "Online status, live scans, emergency open"),
        (400, 250, "User Monitor", "Who is inside campus"),
        (760, 250, "Register Visitor", "Walk-in guest + vehicle"),
        (1120, 250, "Active Visitors", "Assign / return visitor RFID"),
        (40, 450, "Visitor History", "Past visits and RFID returns"),
        (400, 450, "Violations", "Cite plate, photos, 3-strike"),
        (760, 450, "Updates", "Guard notifications"),
        (1120, 450, "Access Logs", "Searchable entry/exit"),
        (40, 650, "Parking", "Live slot occupancy map"),
        (400, 650, "AI Parking Monitor", "YOLO feeds + plate correction"),
        (760, 650, "Plate Lookup", "Owner from plate number"),
        (1120, 650, "Live Cameras", "Camera grid"),
    ]
    for x, y, title, desc in modules:
        rounded(d, (x, y, x + 330, y + 140), WHITE, TEAL)
        d.rectangle((x, y, x + 330, y + 38), fill=TEAL)
        d.text((x + 14, y + 10), title, fill=WHITE, font=font(16, True))
        fnt = font(14)
        for i, line in enumerate(wrap(d, desc, fnt, 300)):
            d.text((x + 14, y + 54 + i * 20), line, fill=GRAY, font=fnt)
    widgets = ["Vehicles inside", "Today's entries", "Active violations", "Available slots",
               "Active / waiting visitors", "RFID still assigned", "Expired visitors",
               "Open Live Gate", "Report Violation", "Recent access activity"]
    for i, w in enumerate(widgets):
        box(d, 40 + (i % 5) * 290, 860 + (i // 5) * 80, 270, 60, w, (204, 251, 241), TEAL)
    return save(img, "fig05_guard_dashboard.png")


def fig_rfid_gate():
    img, d = new_canvas(1400, 1280, "Figure 6. Live Gate Monitor — RFID access flowchart")
    box(d, 480, 110, 440, 70, "Tap RFID on RC522 (Entry or Exit ESP32)", (219, 234, 254), NAVY)
    box(d, 480, 220, 440, 70, "ESP32 POST /api/rfid/scan", NAVY_SOFT)
    box(d, 480, 330, 440, 70, "RfidAccessService.process()", (254, 243, 199), ORANGE)
    dia1 = diamond(d, 700, 490, 320, 110, "UID on user or visitor card?")
    arrow(d, 700, 400, dia1[0], dia1[1])
    nr = box(d, 40, 560, 280, 80, "Card not registered — deny", (254, 226, 226), RED, RED)
    arrow(d, dia1[4], dia1[5], nr[6], nr[5], RED, "No")
    box(d, 500, 640, 400, 120, "Active? Owner+plate? Gate_access? Not already inside/outside?", NAVY_SOFT)
    arrow(d, dia1[2], dia1[3], 700, 640, GREEN, "Yes")
    dia2 = diamond(d, 700, 860, 300, 110, "All checks pass?")
    arrow(d, 700, 760, dia2[0], dia2[1])
    deny = box(d, 40, 940, 300, 90, "Deny — write gate_logs", (254, 226, 226), RED, RED)
    ok = box(d, 980, 820, 360, 100, "granted=true · gate_logs · parking slot sync", (220, 252, 231), GREEN, GREEN)
    arrow(d, dia2[4], dia2[5], deny[6], deny[5], RED, "No")
    arrow(d, dia2[6], dia2[7], ok[4], ok[5], GREEN, "Yes")
    boom = box(d, 500, 1080, 400, 90, "Entry ESP32 servo GPIO 14 opens boom", (204, 251, 241), TEAL)
    arrow(d, 1160, 920, 1160, 1125)
    arrow(d, 1160, 1125, boom[6], boom[5])
    return save(img, "fig06_rfid_gate.png")


def fig_visitor():
    img, d = new_canvas(1400, 900, "Figure 7. Visitor lifecycle (Guard operations)")
    steps = [
        (80, 140, "1. Register Visitor", "Name, purpose, plate, host"),
        (420, 140, "2. Waiting", "Active Visitors list"),
        (760, 140, "3. Assign RFID", "UID bound to this visit"),
        (1100, 140, "4. Gate tap", "Same RFID API as owners"),
        (80, 420, "5. On campus", "status = Inside"),
        (420, 420, "6. Exit tap", "Shared boom opens"),
        (760, 420, "7. Return RFID", "Card freed"),
        (1100, 420, "8. History", "Completed / expired"),
    ]
    for x, y, title, desc in steps:
        rounded(d, (x, y, x + 300, y + 160), WHITE, TEAL)
        d.rectangle((x, y, x + 300, y + 44), fill=TEAL)
        d.text((x + 12, y + 12), title, fill=WHITE, font=font(16, True))
        fnt = font(14)
        for i, line in enumerate(wrap(d, desc, fnt, 270)):
            d.text((x + 14, y + 62 + i * 20), line, fill=GRAY, font=fnt)
    for x in (380, 720, 1060):
        arrow(d, x, 220, x + 40, 220)
    arrow(d, 1250, 300, 1250, 400)
    for x in (1100, 760, 420):
        arrow(d, x, 500, x - 40, 500)
    box(d, 80, 660, 1240, 50, "Collections: visitors, visitor_rfid_cards, gate_logs (visitor_id), parking_slots (parked_visitor_id)", NAVY_SOFT)
    box(d, 80, 740, 1240, 50, "Guard dashboard counts active / waiting / RFID still assigned / expired from these documents.", NAVY_SOFT)
    return save(img, "fig07_visitor_flow.png")


def fig_violations():
    img, d = new_canvas(1400, 1080, "Figure 8. Violations and 3-strike enforcement")
    box(d, 80, 120, 420, 80, "Guard: plate + type + photos", (254, 243, 199), ORANGE)
    box(d, 900, 120, 420, 80, "AI parking auto-cite (optional)", (237, 233, 254), PURPLE)
    box(d, 490, 260, 420, 80, "Match plate → users.plate_number", NAVY_SOFT)
    arrow(d, 290, 200, 490, 290)
    arrow(d, 1110, 200, 910, 290)
    dia = diamond(d, 700, 440, 280, 110, "User found?")
    arrow(d, 700, 340, dia[0], dia[1])
    nf = box(d, 40, 520, 280, 80, "Plate not found", (254, 226, 226), RED, RED)
    arrow(d, dia[4], dia[5], nf[6], nf[5], RED, "No")
    box(d, 500, 560, 400, 80, "Insert violations_log + increment users.strike_count", NAVY_SOFT)
    arrow(d, dia[2], dia[3], 700, 560, GREEN, "Yes")
    box(d, 40, 720, 380, 90, "Strike 1 — warning", (254, 249, 195), GOLD, ORANGE)
    box(d, 510, 720, 380, 90, "Strike 2 — Admin alert list", (255, 237, 213), ORANGE, ORANGE)
    box(d, 980, 720, 380, 90, "Strike 3 — lock user + RFID deny", (254, 226, 226), RED, RED)
    box(d, 300, 880, 800, 90, "May write user_suspensions. Login and gate both blocked at 3 strikes.", (219, 234, 254), NAVY)
    return save(img, "fig08_violations.png")


def fig_ai_parking():
    img, d = new_canvas(1400, 980, "Figure 9. AI parking and plate lookup")
    box(d, 80, 130, 300, 90, "IP camera RTSP", NAVY_SOFT)
    box(d, 480, 130, 420, 90, "YOLOv9 :8090 occupancy + OCR", (237, 233, 254), PURPLE)
    box(d, 1000, 130, 340, 90, "POST /api/ai-parking/*", (219, 234, 254), NAVY)
    arrow(d, 380, 175, 480, 175)
    arrow(d, 900, 175, 1000, 175)
    box(d, 80, 320, 400, 90, "Update parking_slots status", (220, 252, 231), GREEN, GREEN)
    box(d, 560, 320, 360, 90, "Optional auto-citation", (254, 243, 199), ORANGE)
    box(d, 1000, 320, 340, 90, "MJPEG to dashboards", NAVY_SOFT)
    pages = [
        (80, 500, "Admin/Guard Parking", "Slot map from parking_areas + parking_slots"),
        (500, 500, "Guard AI Monitor", "Feeds + correct-plate"),
        (920, 500, "Plate Lookup", "users / visitors by plate"),
        (80, 700, "Admin Live Cameras", "Camera grid"),
        (500, 700, "User Parking", "Allowed zones only"),
        (920, 700, "Student/Staff", "Read-only occupancy"),
    ]
    for x, y, t, desc in pages:
        rounded(d, (x, y, x + 380, y + 130), WHITE, PURPLE)
        d.rectangle((x, y, x + 380, y + 40), fill=PURPLE)
        d.text((x + 12, y + 10), t, fill=WHITE, font=font(15, True))
        fnt = font(14)
        for i, line in enumerate(wrap(d, desc, fnt, 350)):
            d.text((x + 14, y + 56 + i * 20), line, fill=GRAY, font=fnt)
    return save(img, "fig09_ai_parking.png")


def fig_user_map():
    img, d = new_canvas(1400, 980, "Figure 10. Student / Staff dashboard map  |  /user")
    box(d, 420, 110, 560, 90, "User Dashboard — profile, strikes, rules, recent activity", (237, 233, 254), PURPLE)
    modules = [
        (80, 260, "Dashboard", "Vehicle, Gate_access, strikes, campus rules"),
        (500, 260, "Notifications", "Approvals, RFID, citations"),
        (920, 260, "My Violations", "Own citations + evidence"),
        (80, 500, "Entry/Exit History", "Personal gate_logs"),
        (500, 500, "Parking", "Allowed parking_areas only"),
        (920, 500, "Profile", "Update details after approval"),
    ]
    for x, y, t, desc in modules:
        rounded(d, (x, y, x + 380, y + 160), WHITE, PURPLE)
        d.rectangle((x, y, x + 380, y + 42), fill=PURPLE)
        d.text((x + 14, y + 12), t, fill=WHITE, font=font(17, True))
        fnt = font(15)
        for i, line in enumerate(wrap(d, desc, fnt, 350)):
            d.text((x + 16, y + 62 + i * 22), line, fill=GRAY, font=fnt)
    box(d, 80, 720, 1240, 70, "RFID tap at the physical gate writes gate_logs. The website only displays those documents.", NAVY_SOFT)
    box(d, 80, 820, 1240, 70, "If Gate_access is not Granted or strike_count >= 3, the portal may still open but the boom stays closed.", (254, 243, 199), ORANGE)
    return save(img, "fig10_user_dashboard.png")


def fig_er_conceptual():
    img, d = new_canvas(1680, 1180, "Figure 11. Conceptual ER diagram (MongoDB collections)")

    def ent(x, y, w, h, title, color=NAVY):
        rounded(d, (x, y, x + w, y + h), WHITE, color, r=10)
        d.rectangle((x, y, x + w, y + 36), fill=color)
        tw = d.textlength(title, font=font(15, True))
        d.text((x + (w - tw) / 2, y + 8), title, fill=WHITE, font=font(15, True))
        return {"l": x, "r": x + w, "t": y, "b": y + h, "cx": x + w / 2, "cy": y + h / 2, "x": x, "y": y, "w": w, "h": h}

    def diam(cx, cy, text):
        w, h = 150, 70
        pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
        d.polygon(pts, fill=(255, 247, 237), outline=GOLD)
        fnt = font(12, True)
        lines = wrap(d, text, fnt, 120)
        th = len(lines) * 16
        ty = cy - th / 2
        for i, line in enumerate(lines):
            tw = d.textlength(line, font=fnt)
            d.text((cx - tw / 2, ty + i * 16), line, fill=ORANGE, font=fnt)
        return {"l": cx - w / 2, "r": cx + w / 2, "t": cy - h / 2, "b": cy + h / 2, "cx": cx, "cy": cy, "x": cx - w / 2, "y": cy - h / 2, "w": w, "h": h}

    roles = ent(60, 120, 200, 80, "USER_ROLES")
    depts = ent(360, 120, 200, 80, "DEPARTMENTS")
    vehs = ent(660, 120, 200, 80, "VEHICLES")
    vtypes = ent(1420, 120, 200, 80, "VIOLATION_TYPES", ORANGE)
    users = ent(360, 380, 240, 90, "USERS")
    vis = ent(980, 380, 220, 90, "VISITORS", TEAL)
    cards = ent(1320, 380, 260, 90, "VISITOR_RFID_CARDS", TEAL)
    notes = ent(40, 680, 220, 80, "NOTIFICATIONS", PURPLE)
    logs = ent(320, 680, 220, 80, "GATE_LOGS")
    viol = ent(600, 680, 240, 80, "VIOLATIONS_LOG", ORANGE)
    susp = ent(40, 900, 220, 80, "USER_SUSPENSIONS", RED)
    areas = ent(980, 680, 220, 80, "PARKING_AREAS", GREEN)
    slots = ent(980, 900, 220, 80, "PARKING_SLOTS", GREEN)
    rules = ent(1280, 900, 180, 70, "PARKING_RULES", GRAY)
    info = ent(1480, 900, 160, 70, "GEN. INFO", GRAY)
    settings = ent(1280, 680, 200, 80, "SYSTEM_SETTINGS", GRAY)

    d1 = diam(460, 260, "has role")
    d2 = diam(260, 260, "in dept")
    d3 = diam(700, 260, "vehicle type")
    d4 = diam(800, 430, "registered by")
    d5 = diam(1260, 430, "holds card")
    d6 = diam(150, 560, "notified")
    d7 = diam(430, 560, "taps gate")
    d8 = diam(720, 560, "cited")
    d9 = diam(1090, 560, "visits / taps")
    d10 = diam(1090, 820, "contains")
    d11 = diam(700, 820, "occupies")
    d12 = diam(150, 820, "suspended")

    relate(d, roles, "s", d1, "n", "1", "1", "", NAVY)
    relate(d, d1, "s", users, "n", "1", "N", "1 : N", NAVY)
    relate(d, depts, "s", d2, "e", "1", "1", "", NAVY)
    relate(d, d2, "s", users, "n", "1", "N", "", NAVY)
    relate(d, vehs, "s", d3, "n", "1", "1", "", NAVY)
    relate(d, d3, "s", users, "e", "1", "N", "", NAVY)
    relate(d, d3, "e", vis, "n", "1", "N", "", TEAL)
    relate(d, users, "e", d4, "w", "1", "1", "", TEAL)
    relate(d, d4, "e", vis, "w", "1", "N", "registers", TEAL)
    relate(d, vis, "e", d5, "w", "1", "1", "", TEAL)
    relate(d, d5, "e", cards, "w", "1", "0..1", "0..1", TEAL)
    relate(d, users, "s", d6, "n", "1", "1", "", PURPLE)
    relate(d, d6, "s", notes, "n", "1", "N", "", PURPLE)
    relate(d, users, "s", d7, "n", "1", "1", "", NAVY)
    relate(d, vis, "s", d9, "n", "1", "1", "", TEAL)
    relate(d, d7, "s", logs, "n", "1", "N", "", NAVY)
    relate(d, d9, "s", logs, "e", "1", "N", "", TEAL)
    relate(d, users, "s", d8, "n", "1", "1", "", ORANGE)
    relate(d, d8, "s", viol, "n", "1", "N", "", ORANGE)
    relate(d, vtypes, "s", viol, "e", "1", "N", "by name", ORANGE)
    relate(d, users, "s", d12, "n", "1", "1", "", RED)
    relate(d, d12, "s", susp, "n", "1", "0..1", "0..1", RED)
    relate(d, areas, "s", d10, "n", "1", "1", "", GREEN)
    relate(d, d10, "s", slots, "n", "1", "N", "1 : N", GREEN)
    relate(d, users, "s", d11, "w", "1", "1", "", GREEN)
    relate(d, vis, "s", d11, "e", "1", "1", "", GREEN)
    relate(d, d11, "s", slots, "n", "0..1", "0..1", "optional", GREEN)

    d.text((40, 1100), "Chen notation: rectangles = entities (collections), diamonds = relationships. Crow’s-foot detail is Figure 12.", fill=GRAY, font=font(14))
    d.text((40, 1135), "Logical FKs are integer id fields (not MongoDB ObjectId). violation_types is linked by violation_name text, not a numeric FK.", fill=GRAY, font=font(14))
    return save(img, "fig11_er_conceptual.png")


def fig_er_logical():
    img, d = new_canvas(2000, 1680, "Figure 12. Logical ER diagram — collections, keys, and cardinalities")

    roles = entity(d, 40, 110, 230, "user_roles", [("PK", "id"), ("A", "role_name")])
    depts = entity(d, 40, 280, 230, "departments", [("PK", "departmentcode"), ("A", "departmentname")])
    vehs = entity(d, 40, 450, 230, "vehicles", [("PK", "id"), ("A", "vehicle_name")])
    users = entity(d, 360, 200, 280, "users", [
        ("PK", "id"), ("A", "fullname"), ("A", "email"), ("A", "password"),
        ("A", "id_number"), ("A", "plate_number"), ("A", "rfid_uid"),
        ("A", "status"), ("A", "Gate_access"), ("A", "strike_count"),
        ("FK", "user_role_id"), ("FK", "department_code"), ("FK", "vehicle_id"),
        ("A", "driver_license"), ("A", "or_cr_photo"),
    ])
    uvehs = entity(d, 360, 520, 280, "user_vehicles", [
        ("PK", "id"), ("FK", "user_id"), ("FK", "vehicle_id"),
        ("A", "plate_number"), ("A", "vehicle_model"), ("A", "is_primary"),
    ])
    vis = entity(d, 760, 110, 270, "visitors", [
        ("PK", "id"), ("A", "first_name / last_name"), ("A", "purpose"),
        ("A", "plate_number"), ("A", "status"), ("A", "time_in / time_out"),
        ("FK", "vehicle_id"), ("FK", "registered_by"),
        ("FK", "visitor_rfid_card_id"), ("A", "rfid_uid"),
    ])
    cards = entity(d, 1120, 110, 270, "visitor_rfid_cards", [
        ("PK", "id"), ("A", "rfid_uid"), ("A", "status"),
        ("FK", "visitor_id"), ("FK", "created_by"),
        ("A", "assigned_at"), ("A", "returned_at"), ("A", "expires_at"),
    ])
    logs = entity(d, 760, 680, 270, "gate_logs", [
        ("PK", "id"), ("FK", "user_id"), ("FK", "visitor_id"),
        ("A", "action  Entry|Exit"), ("A", "gate_id"), ("A", "rfid_uid"),
        ("A", "result"), ("A", "reason"), ("A", "timestamp"),
    ])
    viol = entity(d, 1120, 430, 280, "violations_log", [
        ("PK", "id"), ("FK", "user_id"), ("A", "guard_id"),
        ("A", "violator_name"), ("A", "plate_number"),
        ("A", "violation_type"), ("A", "evidence_photos[]"),
        ("FK", "area_id"), ("A", "status"), ("A", "camera_id"),
    ])
    notes = entity(d, 360, 780, 280, "notifications", [
        ("PK", "id"), ("FK", "user_id"), ("FK", "sender_id"),
        ("A", "title / message"), ("A", "type"),
        ("FK", "violation_log_id"), ("A", "is_read"),
    ])
    susp = entity(d, 40, 780, 230, "user_suspensions", [
        ("PK", "id"), ("FK", "user_id"), ("A", "strike_count"),
        ("A", "is_suspended"), ("A", "suspended_until"),
    ])
    areas = entity(d, 1480, 110, 270, "parking_areas", [
        ("PK", "id"), ("A", "area_name"), ("A", "capacity"),
        ("A", "allowed_roles[]"), ("A", "is_visible"),
        ("A", "designation_notes"),
    ])
    slots = entity(d, 1480, 380, 270, "parking_slots", [
        ("PK", "id"), ("FK", "area_id"), ("A", "slot_number"),
        ("A", "status"), ("FK", "parked_user_id"),
        ("FK", "parked_visitor_id"),
    ])
    vtypes = entity(d, 1120, 820, 270, "violation_types", [
        ("PK", "id"), ("A", "violation_name"), ("A", "description"), ("A", "status"),
    ])
    settings = entity(d, 1480, 700, 270, "system_settings", [
        ("PK", "id = 1"), ("A", "campus_name"),
        ("A", "auto_lock_on_3rd_violation"),
        ("A", "require_photo_evidence"),
        ("A", "enable_visitor_time_limits"),
    ])
    rules = entity(d, 760, 900, 250, "parking_rules", [("PK", "id"), ("A", "description")])
    ginfo = entity(d, 400, 1080, 250, "general_informations", [("PK", "id"), ("A", "description")])
    rperm = entity(d, 40, 1080, 250, "role_permissions", [("PK", "id = 1"), ("A", "matrix{}")])
    sanc = entity(d, 1120, 1080, 250, "violation_sanctions", [("PK", "id"), ("A", "sanctions_name")])
    stall = entity(d, 1480, 1080, 250, "stalled_vehicles", [("PK", "id"), ("A", "description")])

    relate(d, roles, "e", users, "w", "1", "N", "user_role_id")
    relate(d, depts, "e", users, "w", "1", "N", "department_code")
    relate(d, vehs, "e", users, "w", "1", "N", "vehicle_id")
    relate(d, users, "s", uvehs, "n", "1", "N", "user_id", NAVY)
    relate(d, vehs, "e", uvehs, "w", "1", "N", "vehicle_id", NAVY)
    relate(d, vehs, "e", vis, "w", "1", "N", "vehicle_id", TEAL)
    relate(d, users, "e", vis, "w", "1", "N", "registered_by", TEAL)
    relate(d, vis, "e", cards, "w", "1", "0..1", "card_id / visitor_id", TEAL)
    relate(d, users, "e", cards, "s", "1", "N", "created_by", TEAL)
    relate(d, users, "e", logs, "w", "1", "0..N", "user_id")
    relate(d, vis, "s", logs, "n", "1", "0..N", "visitor_id", TEAL)
    relate(d, users, "e", viol, "w", "1", "0..N", "user_id", ORANGE)
    relate(d, areas, "s", viol, "e", "1", "0..N", "area_id", ORANGE)
    relate(d, users, "s", notes, "n", "1", "N", "user_id", PURPLE)
    relate(d, viol, "s", notes, "e", "1", "0..N", "violation_log_id", PURPLE)
    relate(d, users, "w", susp, "e", "1", "0..1", "user_id", RED)
    relate(d, areas, "s", slots, "n", "1", "N", "area_id", GREEN)
    relate(d, users, "e", slots, "w", "1", "0..1", "parked_user_id", GREEN)
    relate(d, vis, "e", slots, "w", "1", "0..1", "parked_visitor_id", GREEN)

    # legend
    d.rounded_rectangle((40, 1420, 1960, 1620), radius=10, fill=WHITE, outline=NAVY, width=2)
    d.text((60, 1436), "Legend and notes", fill=NAVY, font=font(16, True))
    d.rectangle((60, 1474, 90, 1496), fill=PK_BG, outline=NAVY)
    d.text((100, 1474), "PK  integer id (HasSequentialId) — used like a SQL primary key", fill=NAVY, font=font(14))
    d.rectangle((60, 1506, 90, 1528), fill=FK_BG, outline=NAVY)
    d.text((100, 1506), "FK  integer reference stored on the child document (MongoDB has no enforced foreign keys)", fill=NAVY, font=font(14))
    d.text((60, 1544), "Crow’s foot  N / 0..N = many     |     bar = one     |     circle = optional (0..1).", fill=NAVY, font=font(14))
    d.text((60, 1574), "violations_log.violation_type stores the name from violation_types (logical link, not numeric FK). "
           "role_permissions and system_settings are singleton documents (id = 1).", fill=GRAY, font=font(14))

    return save(img, "fig12_er_logical.png")


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*NAVY)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_img(doc, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run = cap.add_run(caption)
    set_run_font(run, 10, True, NAVY)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, 10, True, WHITE)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0F2D5A")
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 10)
    doc.add_paragraph()


def build_docx(images: dict[str, Path]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(72)
    r = t.add_run("CAMARINES SUR POLYTECHNIC COLLEGES")
    set_run_font(r, 14, True, NAVY)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run("Smart Campus Vehicle Management System")
    set_run_font(r, 22, True, NAVY)
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t3.add_run("Dashboards, Flowcharts, and Database ER Diagram")
    set_run_font(r, 14, False, GOLD)
    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t4.paragraph_format.space_before = Pt(18)
    r = t4.add_run("Laravel 12 + MongoDB  ·  ESP32 RFID  ·  YOLOv9 parking")
    set_run_font(r, 11, False, GRAY)

    doc.add_page_break()
    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. System overview",
        "2. Database entity-relationship diagram",
        "3. Login and role-based routing",
        "4. Admin portal",
        "5. Guard portal",
        "6. Student / Staff portal",
        "7. Cross-cutting process flows",
        "8. URL summary by role",
    ]:
        add_p(doc, item)

    add_heading(doc, "1. System overview", 1)
    add_p(doc, "The Smart Campus Vehicle Management System (VMS) controls campus vehicle entry, visitor RFID, parking occupancy, and a three-strike violation policy. The website is Laravel 12 with MongoDB. ESP32 boards at Entry and Exit call the RFID API; Laravel decides grant or deny. Optional YOLOv9 cameras update parking slots.")
    add_p(doc, "Three signed-in dashboards exist: Admin (/admin), Guard (/guard), and Student/Staff (/user). Visitors have no website login; Guards register them and issue a temporary RFID card.")
    add_img(doc, images["arch"], "Figure 1. System architecture — hardware, APIs, dashboards, and MongoDB.")

    add_heading(doc, "2. Database entity-relationship diagram", 1)
    add_p(doc, "Data lives in MongoDB (database name capstone). Each Eloquent model maps to a collection. Documents use a sequential integer id (HasSequentialId) as a logical primary key, and child documents store integer foreign keys. MongoDB does not enforce those FKs; Laravel relationships do.")
    add_p(doc, "Figure 11 is the conceptual (Chen) view used in oral defense. Figure 12 is the logical Crow’s-foot diagram with attributes taken from the model $fillable arrays.")
    add_img(doc, images["er_c"], "Figure 11. Conceptual ER diagram of the main collections and relationships.")
    add_img(doc, images["er_l"], "Figure 12. Logical ER diagram with primary keys, foreign keys, and cardinalities.")

    add_heading(doc, "2.1 Collection catalog", 2)
    add_table(doc, ["Collection", "Model", "Role in the system"], [
        ["users", "User", "Admin, Guard, Student, Staff accounts; plate, RFID UID, strikes, Gate_access"],
        ["user_vehicles", "UserVehicle", "Multiple vehicles per user (plate, model, primary flag)"],
        ["user_roles", "UserRole", "Admin, Guard, Student, Staff, Visitor role names"],
        ["departments", "Department", "College/office codes for staff/students"],
        ["vehicles", "Vehicle", "Lookup of vehicle types (car, motorcycle, …)"],
        ["visitors", "Visitor", "Walk-in guests registered by Guard"],
        ["visitor_rfid_cards", "VisitorRfidCard", "Temporary RFID pool assigned to a visit"],
        ["gate_logs", "GateLog", "Every RFID tap: Entry/Exit, granted or denied"],
        ["parking_areas", "ParkingArea", "Lots/zones and allowed_roles"],
        ["parking_slots", "ParkingSlot", "Slot status; optional parked user or visitor"],
        ["violations_log", "ViolationLog", "Citations + evidence; drives strike_count"],
        ["violation_types", "ViolationType", "Catalog of citation names (Admin Settings)"],
        ["notifications", "Notification", "In-app messages to users/guards"],
        ["user_suspensions", "UserSuspension", "Lock record after 3 strikes"],
        ["parking_rules", "ParkingRule", "Official rules shown on /user dashboard"],
        ["general_informations", "GeneralInformation", "Campus notices on /user dashboard"],
        ["system_settings", "SystemSetting", "Singleton: auto-lock, photo evidence, visitor time limits"],
        ["role_permissions", "RolePermission", "Singleton permission matrix"],
        ["violation_sanctions", "ViolationSanction", "Sanction name lookup"],
        ["stalled_vehicles", "StalledVehicle", "Stalled-vehicle note lookup"],
    ])
    add_p(doc, "Core cardinalities: one user_role has many users; one user has many gate_logs, notifications, and violations_log; one parking_area has many parking_slots; a visitor optionally holds one visitor_rfid_card (and the card points back via visitor_id). A parking_slot is occupied by at most one user or one visitor. A gate_log belongs to either a user_id or a visitor_id, not both.")

    add_heading(doc, "3. Login and role-based routing", 1)
    add_p(doc, "After password check, email must be verified. Pending, denied, or 3-strike locked accounts cannot open a dashboard. Granted users go to Admin → /admin, Guard → /guard, Student or Staff → /user.")
    add_img(doc, images["login"], "Figure 2. Authentication and dashboard routing flowchart.")

    add_heading(doc, "4. Admin portal", 1)
    add_heading(doc, "4.1 Admin Dashboard (/admin)", 2)
    add_p(doc, "Campus health check: user counts, active violations, today’s gate_logs, parking occupancy, visitor counts, weekly trends, violation-type chart, 2nd-strike watch list, and recent violations_log rows. It does not open the boom.")
    add_img(doc, images["admin"], "Figure 3. Admin dashboard widgets and module map.")
    add_heading(doc, "4.2 Registrations (/admin/registrations)", 2)
    add_p(doc, "Pending users become Granted or Denied. Approval opens /user. RFID gate access stays off until RFID Assignment. Guards are created under Settings / Create Guard.")
    add_heading(doc, "4.3 RFID Assignment (/admin/rfid)", 2)
    add_p(doc, "Binds users.rfid_uid and sets Gate_access = Granted so ESP32 taps can succeed.")
    add_img(doc, images["reg"], "Figure 4. From public registration to RFID-enabled gate access.")
    add_heading(doc, "4.4 User Management, visitors, violations, logs", 2)
    add_p(doc, "User Management audits documents and strikes. Admin visitor pages are oversight only. Violations are read-only evidence. Access Logs are the searchable gate_logs feed (Reverb live).")
    add_heading(doc, "4.5 Parking, cameras, reports, settings", 2)
    add_p(doc, "Parking edits parking_areas / parking_slots and zone allowed_roles. Live Cameras show AI streams. Reports export analytics. Settings write parking_rules, general_informations, violation_types, and system_settings.")

    add_heading(doc, "5. Guard portal", 1)
    add_heading(doc, "5.1 Guard Dashboard (/guard)", 2)
    add_p(doc, "Duty-officer home: vehicles inside, today’s entries, active violations, free slots, visitor RFID pool, shortcuts to Live Gate and Report Violation, recent gate_logs and citations.")
    add_img(doc, images["guard"], "Figure 5. Guard dashboard widgets and module map.")
    add_heading(doc, "5.2 Live Gate Monitor (/guard/gate)", 2)
    add_p(doc, "ESP32 heartbeat = online. Scans arrive on Reverb gate.scans. Emergency Open queues the Entry servo (GPIO 14). Exit grants open the same shared boom.")
    add_img(doc, images["rfid"], "Figure 6. RFID tap decision flowchart.")
    add_heading(doc, "5.3 Visitors", 2)
    add_p(doc, "Register Visitor → assign visitor_rfid_cards → gate tap writes gate_logs.visitor_id → return card → history.")
    add_img(doc, images["visitor"], "Figure 7. Visitor lifecycle.")
    add_heading(doc, "5.4 Violations, AI parking, plate lookup", 2)
    add_p(doc, "Citations insert violations_log and increment users.strike_count. AI Parking Monitor corrects plates and may auto-cite. Plate Lookup resolves users or visitors by plate.")

    add_heading(doc, "6. Student / Staff portal", 1)
    add_p(doc, "Personal portal only. Dashboard shows plate, Gate_access, strike_count, parking_rules, general_informations, recent own gate_logs and violations. Notifications, My Violations, Entry/Exit History, and Parking (allowed zones) are read-oriented. This role cannot open the boom or cite others.")
    add_img(doc, images["user"], "Figure 10. Student/Staff dashboard and modules.")

    add_heading(doc, "7. Cross-cutting process flows", 1)
    add_p(doc, "RFID: Entry ESP32 has the servo; Exit has RFID only. Fail-closed on network errors. Three-strike: strike 2 alerts Admin; strike 3 locks login and RFID. AI parking posts occupancy into parking_slots and optional events into violations_log.")
    add_img(doc, images["viol"], "Figure 8. Violation recording and strike escalation.")
    add_img(doc, images["ai"], "Figure 9. Camera → YOLOv9 → Laravel occupancy → dashboards.")

    add_heading(doc, "8. URL summary by role", 1)
    add_table(doc, ["Role", "URL", "Page"], [
        ["Admin", "/admin", "Dashboard"],
        ["Admin", "/admin/registrations", "Registrations"],
        ["Admin", "/admin/rfid", "RFID Assignment"],
        ["Admin", "/admin/users", "User Management"],
        ["Admin", "/admin/parking", "Parking + zones"],
        ["Admin", "/admin/reports", "Reports"],
        ["Admin", "/admin/settings", "Settings"],
        ["Guard", "/guard", "Dashboard"],
        ["Guard", "/guard/gate", "Live Gate Monitor"],
        ["Guard", "/guard/visitors/register", "Register Visitor"],
        ["Guard", "/guard/violations", "Violations"],
        ["Guard", "/guard/ai-parking", "AI Parking Monitor"],
        ["Guard", "/guard/plate-lookup", "Plate Lookup"],
        ["Student/Staff", "/user", "Dashboard"],
        ["Student/Staff", "/user/notifications", "Notifications"],
        ["Student/Staff", "/user/violations", "My Violations"],
        ["Student/Staff", "/user/entry-exit", "Entry/Exit History"],
        ["Student/Staff", "/user/parking", "Parking"],
    ])
    add_p(doc, "APIs: POST /api/rfid/scan, POST /api/rfid/heartbeat, POST /api/ai-parking/occupancy, POST /api/ai-parking/events, POST /api/ai-parking/plate-lookup.")

    footer = doc.add_paragraph()
    r = footer.add_run("Generated from Laravel models and routes. MongoDB does not enforce FKs; integer ids are application-level keys.")
    set_run_font(r, 9, False, GRAY)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main() -> None:
    images = {
        "arch": fig_architecture(),
        "login": fig_login(),
        "admin": fig_admin_map(),
        "reg": fig_registration(),
        "guard": fig_guard_map(),
        "rfid": fig_rfid_gate(),
        "visitor": fig_visitor(),
        "viol": fig_violations(),
        "ai": fig_ai_parking(),
        "user": fig_user_map(),
        "er_c": fig_er_conceptual(),
        "er_l": fig_er_logical(),
    }
    path = build_docx(images)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
